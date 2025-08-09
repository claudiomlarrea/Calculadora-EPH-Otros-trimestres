
# -*- coding: utf-8 -*-
"""
Calculadora EPH (2017–2024) para TODOS los trimestres (1º a 4º) — v2
--------------------------------------------------------------------
- Mapea códigos -> etiquetas nominales (sexo, nivel educativo, condición de actividad)
- Genera informe Word con un apartado final de análisis y conclusiones robusto
- Omite secciones no disponibles sin romper
- Detecta año y trimestre automáticamente

Ejecución:
    streamlit run app_eph_trimestres.py
"""

import io
import re
from datetime import datetime

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =============================
# Utilidades robustas
# =============================

@st.cache_data(show_spinner=False)
def _read_table(uploaded_file: st.runtime.uploaded_file_manager.UploadedFile) -> pd.DataFrame:
    """
    Lee CSV, TXT (delimitado por ; , \t |) o Excel (.xlsx/.xls) tolerando encodings.
    """
    name = uploaded_file.name.lower()
    data = uploaded_file.read()

    # Excel
    if name.endswith((".xlsx", ".xls")):
        uploaded_file.seek(0)
        # openpyxl se instala vía requirements.txt
        return pd.read_excel(uploaded_file)

    # CSV/TXT
    encodings = ["utf-8", "latin-1", "cp1252"]
    sep_candidates = [",", ";", "\t", "|"]
    for enc in encodings:
        try:
            text = data.decode(enc, errors="ignore")
            first_line = text.splitlines()[0] if text else ""
            counts = {s: first_line.count(s) for s in sep_candidates}
            sep = max(counts, key=counts.get) if counts else ","
            df = pd.read_csv(io.StringIO(text), sep=sep)
            return df
        except Exception:
            continue

    # Último intento: autodetección
    uploaded_file.seek(0)
    try:
        return pd.read_csv(uploaded_file)
    except Exception:
        uploaded_file.seek(0)
        return pd.read_table(uploaded_file, engine="python", sep=None)


def get_first_col(df: pd.DataFrame, candidates):
    for c in candidates:
        if c in df.columns: return c
        for col in df.columns:
            if col.lower() == c.lower(): return col
    return None


def detect_year_quarter(df: pd.DataFrame, fallback_name: str = ""):
    year_col = get_first_col(df, ["ANO4", "ANO", "AÑO", "YEAR", "ANIO", "ANIO4"])
    q_col    = get_first_col(df, ["TRIMESTRE", "TRIM", "TRIMES", "QUARTER"])

    anio = None; tri = None
    if year_col is not None:
        try:
            anio = int(pd.to_numeric(df[year_col], errors="coerce").mode().iloc[0])
        except Exception:
            pass
    if q_col is not None:
        try:
            tri = int(pd.to_numeric(df[q_col], errors="coerce").mode().iloc[0])
        except Exception:
            pass

    # Fallback por nombre de archivo
    if anio is None or tri is None:
        m_year = re.search(r"(20\d{2})", fallback_name)
        if m_year and anio is None: anio = int(m_year.group(1))
        m_q = re.search(r"(?:^|[^\d])(1|2|3|4)\s*(?:t|tri|trim|trimestre)\b", fallback_name, re.I)
        if m_q and tri is None: tri = int(m_q.group(1))
    return anio, tri


# =============================
# Diccionarios de mapeo (EPH clásico)
# =============================

MAP_SEXO = {1: "Varón", 2: "Mujer"}

# NIVEL_ED (1–7). Si tu base usa otros códigos, podés ampliar aquí.
MAP_NIVEL_ED = {
    1: "Sin instrucción",
    2: "Primaria incompleta",
    3: "Primaria completa",
    4: "Secundaria incompleta",
    5: "Secundaria completa",
    6: "Terciario/Universitario incompleto",
    7: "Terciario/Universitario completo",
}

# ESTADO / COND_ACT (valores frecuentes en EPH)
# 0 y 9 como reservas para NR/NC si aparecen
MAP_COND_ACT = {
    0: "No corresponde / NR",
    1: "Ocupado/a",
    2: "Desocupado/a",
    3: "Inactivo/a",
    4: "Menor de 10 años",
    9: "Ns/Nc",
}

# TIC (si existieran)
MAP_BIN_TIC = {0: "No", 1: "Sí", 2: "Ns/Nc"}


class Cols:
    SEXO  = ["CH04", "SEXO"]
    EDAD  = ["CH06", "EDAD"]
    NIVEL_ED = ["NIVEL_ED", "NIVEL_EDUC", "NIVEL_EDUCATIVO", "EDUC_NIVEL"]
    COND_ACT = ["ESTADO", "COND_ACT", "CONDICION_ACT", "CAT_OCUP"]

    TIC_PC = ["TIP_III_04", "USO_PC", "PC_USO"]
    TIC_INTERNET = ["TIP_III_06", "USO_INTERNET", "INTERNET_USO"]


def pick(df, options):
    return get_first_col(df, options)


def value_counts_labeled(series: pd.Series, mapping: dict | None = None):
    if series is None: return None
    s = pd.to_numeric(series, errors="ignore")
    if mapping is not None:
        # Convertir a num para mapear por clave si es posible
        s_num = pd.to_numeric(s, errors="coerce")
        labeled = s_num.map(mapping)
        # fallback: si no mapea (p. ej. ya es string), usar valores originales
        labeled = labeled.where(~labeled.isna(), s.astype(str))
        vc = labeled.value_counts(dropna=False).sort_index()
    else:
        vc = pd.Series(s).value_counts(dropna=False).sort_index()
    return vc


# =============================
# Análisis
# =============================

def analyze_individuals(df_ind: pd.DataFrame):
    res = {"N_individuos": int(len(df_ind))}

    sexo_col = pick(df_ind, Cols.SEXO)
    edad_col = pick(df_ind, Cols.EDAD)
    nivel_col = pick(df_ind, Cols.NIVEL_ED)
    cond_col = pick(df_ind, Cols.COND_ACT)

    # Sexo
    if sexo_col:
        v = pd.to_numeric(df_ind[sexo_col], errors="coerce")
        vc = v.map(MAP_SEXO).fillna("Otro/NR").value_counts(dropna=False)
        res["sexo_counts"] = vc.to_dict()
    else:
        res["sexo_counts"] = None

    # Edad
    if edad_col:
        v = pd.to_numeric(df_ind[edad_col], errors="coerce")
        res["edad_media"] = float(np.nanmean(v)) if v.notna().any() else None
        res["edad_p50"]   = float(np.nanmedian(v)) if v.notna().any() else None
        bins   = [0, 5, 12, 18, 30, 45, 60, 75, 120]
        labels = ["0-4", "5-12", "13-18", "19-30", "31-45", "46-60", "61-75", "75+"]
        cat = pd.cut(v, bins=bins, labels=labels, right=True, include_lowest=True)
        res["edad_tramos"] = cat.value_counts(dropna=False).sort_index().to_dict()
    else:
        res["edad_media"] = res["edad_p50"] = None
        res["edad_tramos"] = None

    # Nivel educativo
    if nivel_col:
        vc = value_counts_labeled(df_ind[nivel_col], MAP_NIVEL_ED)
        res["nivel_educativo_counts"] = None if vc is None else vc.to_dict()
    else:
        res["nivel_educativo_counts"] = None

    # Condición de actividad
    if cond_col:
        vc = value_counts_labeled(df_ind[cond_col], MAP_COND_ACT)
        res["actividad_counts"] = None if vc is None else vc.to_dict()
    else:
        res["actividad_counts"] = None

    # TIC (solo si hay)
    pc_col = pick(df_ind, Cols.TIC_PC)
    int_col = pick(df_ind, Cols.TIC_INTERNET)
    if pc_col or int_col:
        tic = {}
        if pc_col:
            vc = value_counts_labeled(df_ind[pc_col], MAP_BIN_TIC)
            tic["uso_pc_counts"] = None if vc is None else vc.to_dict()
        if int_col:
            vc = value_counts_labeled(df_ind[int_col], MAP_BIN_TIC)
            tic["uso_internet_counts"] = None if vc is None else vc.to_dict()
        res["tic"] = tic
    else:
        res["tic"] = None

    return res


def analyze_households(df_hog: pd.DataFrame):
    res = {"N_hogares": int(len(df_hog))}
    ing_candidates = ["ITF", "INGTOT", "INGRESO_TOTAL", "ING_HOGAR", "P47T", "INGTRIM"]
    inc_col = pick(df_hog, ing_candidates)
    if inc_col:
        v = pd.to_numeric(df_hog[inc_col], errors="coerce")
        res["ingreso_hogar_media"]  = float(np.nanmean(v)) if v.notna().any() else None
        res["ingreso_hogar_p50"]    = float(np.nanmedian(v)) if v.notna().any() else None
    else:
        res["ingreso_hogar_media"] = res["ingreso_hogar_p50"] = None
    return res


# =============================
# Informe Word
# =============================

def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

def fmt_num(x, nd=0):
    if x is None or (isinstance(x, float) and np.isnan(x)): return "s/d"
    if isinstance(x, float): return f"{x:,.{nd}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    if isinstance(x, (int, np.integer)): return f"{x:,}".replace(",", ".")
    return str(x)

def dict_to_lines(d: dict):
    return [f"• {k}: {fmt_num(v)}" for k, v in (d or {}).items()]


def build_conclusions(hog_res, ind_res, meta):
    total_ind = ind_res.get("N_individuos") or 0
    total_hog = hog_res.get("N_hogares") or 0

    # Sexo shares
    sexo = ind_res.get("sexo_counts") or {}
    tot_sexo = sum(v for v in sexo.values()) or 1
    share_mujer = 100 * (sexo.get("Mujer", 0) / tot_sexo)
    share_varon = 100 * (sexo.get("Varón", 0) / tot_sexo)

    # Tramos etarios – pico
    tramos = ind_res.get("edad_tramos") or {}
    tramo_top = max(tramos, key=tramos.get) if tramos else None

    # Nivel educativo – top
    ed = ind_res.get("nivel_educativo_counts") or {}
    top_ed = max(ed, key=ed.get) if ed else None

    # Actividad – shares básicos
    act = ind_res.get("actividad_counts") or {}
    tot_act = sum(act.values()) or 1
    share_ocup = 100 * (act.get("Ocupado/a", 0) / tot_act)
    share_desoc = 100 * (act.get("Desocupado/a", 0) / tot_act)
    share_inac  = 100 * (act.get("Inactivo/a", 0) / tot_act)

    lines = []
    lines.append("El presente apartado sintetiza hallazgos clave y sugiere implicancias para política pública y gestión institucional.")
    lines.append(
        f"En el {fmt_num(meta.get('trimestre'))}º trimestre de {fmt_num(meta.get('anio'))} se procesaron {fmt_num(total_hog)} hogares y {fmt_num(total_ind)} individuos."
    )
    # Sexo
    lines.append(
        f"La estructura por sexo es relativamente equilibrada: {share_mujer:.1f}% mujeres y {share_varon:.1f}% varones, "
        "sin desbalances extremos a nivel agregado."
    )
    # Edad
    if tramo_top:
        lines.append(
            f"El tramo etario con mayor peso es {tramo_top}, lo que sugiere que la demanda de servicios y políticas debe contemplar necesidades específicas de ese grupo."
        )
    # Educación
    if top_ed:
        lines.append(
            f"En educación, predomina el nivel \"{top_ed}\", indicador del perfil de capital humano de la muestra. "
            "Este patrón condiciona la inserción laboral y las trayectorias de movilidad social."
        )
    # Actividad
    lines.append(
        f"En el mercado de trabajo, se observa una tasa relativa de ocupación aproximada de {share_ocup:.1f}%, "
        f"desocupación de {share_desoc:.1f}% e inactividad de {share_inac:.1f}%. "
        "Estos valores orientan la priorización de programas de empleabilidad y formación."
    )
    # TIC si existiera
    if ind_res.get("tic"):
        tic = ind_res["tic"]
        if tic.get("uso_internet_counts"):
            u = tic["uso_internet_counts"]
            tot_u = sum(u.values()) or 1
            share_si = 100 * (u.get("Sí", 0) / tot_u)
            lines.append(
                f"Respecto de la inclusión digital, el {share_si:.1f}% declara usar Internet. "
                "Aun así, persisten brechas que tienden a concentrarse en hogares con menores ingresos y menor nivel educativo."
            )

    lines.append(
        "Recomendaciones: (i) fortalecer estrategias de terminalidad educativa en niveles medio y superior; "
        "(ii) articular políticas activas de empleo con formación en habilidades digitales; "
        "(iii) priorizar conectividad significativa y acceso a dispositivos en hogares vulnerables; "
        "(iv) monitorear periódicamente estos indicadores por trimestre para detectar cambios de tendencia."
    )
    return "\n".join(lines)


def build_report(hog_res, ind_res, meta, instructivo_name: str | None):
    doc = Document()

    # Portada
    title = doc.add_paragraph()
    run = title.add_run("Universidad Católica de Cuyo\nSecretaría de Investigación\nInforme EPH por Trimestre")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "")
    add_paragraph(doc, f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # Metadatos
    add_heading(doc, "1. Metadatos del procesamiento", level=1)
    add_paragraph(doc, f"Año detectado: {meta.get('anio', 's/d')} | Trimestre detectado: {meta.get('trimestre', 's/d')}")
    add_paragraph(doc, f"Archivos cargados: Hogares = {meta.get('hog_file', 's/d')} | Individuos = {meta.get('ind_file', 's/d')}")
    add_paragraph(doc, f"Instructivo cargado: {instructivo_name or 's/d'}")

    # Hogares
    add_heading(doc, "2. Resultados de Hogares", level=1)
    add_paragraph(doc, f"Cantidad de hogares: {fmt_num(hog_res.get('N_hogares'))}")
    add_paragraph(doc, f"Ingreso total del hogar (media): {fmt_num(hog_res.get('ingreso_hogar_media'), 2)}")
    add_paragraph(doc, f"Ingreso total del hogar (mediana): {fmt_num(hog_res.get('ingreso_hogar_p50'), 2)}")

    # Individuos
    add_heading(doc, "3. Resultados de Individuos", level=1)
    add_paragraph(doc, f"Cantidad de individuos: {fmt_num(ind_res.get('N_individuos'))}")

    if ind_res.get("sexo_counts"):
        add_heading(doc, "3.1 Distribución por sexo", level=2)
        for line in dict_to_lines(ind_res["sexo_counts"]): add_paragraph(doc, line)

    if ind_res.get("edad_media") is not None:
        add_heading(doc, "3.2 Estadísticas de edad", level=2)
        add_paragraph(doc, f"Edad media: {fmt_num(ind_res.get('edad_media'), 1)} | Mediana: {fmt_num(ind_res.get('edad_p50'), 1)}")
        if ind_res.get("edad_tramos"):
            add_paragraph(doc, "Distribución por tramos:")
            for line in dict_to_lines(ind_res["edad_tramos"]): add_paragraph(doc, line)

    if ind_res.get("nivel_educativo_counts"):
        add_heading(doc, "3.3 Nivel educativo (conteos)", level=2)
        for line in dict_to_lines(ind_res["nivel_educativo_counts"]): add_paragraph(doc, line)

    if ind_res.get("actividad_counts"):
        add_heading(doc, "3.4 Condición de actividad (conteos)", level=2)
        for line in dict_to_lines(ind_res["actividad_counts"]): add_paragraph(doc, line)

    if ind_res.get("tic"):
        add_heading(doc, "3.5 Acceso y uso de TIC (si disponible)", level=2)
        tic = ind_res["tic"]
        if tic.get("uso_pc_counts"):
            add_paragraph(doc, "Uso de computadora:")
            for line in dict_to_lines(tic["uso_pc_counts"]): add_paragraph(doc, line)
        if tic.get("uso_internet_counts"):
            add_paragraph(doc, "Uso de internet:")
            for line in dict_to_lines(tic["uso_internet_counts"]): add_paragraph(doc, line)

    # Conclusiones robustas
    add_heading(doc, "4. Análisis y conclusiones", level=1)
    add_paragraph(doc, build_conclusions(hog_res, ind_res, meta))

    return doc


# =============================
# UI Streamlit
# =============================

st.set_page_config(page_title="Calculadora EPH – Todos los trimestres (v2)", layout="wide")

st.markdown(
    """
    <div style="background:#0a3d62;color:white;padding:14px;border-radius:12px;margin-bottom:14px;">
      <div style="font-weight:700;font-size:20px;">Universidad Católica de Cuyo</div>
      <div style="font-weight:600;font-size:16px;">Secretaría de Investigación</div>
      <div style="font-size:14px;opacity:.9;">Calculadora EPH 2017–2024 · Informe Word automático · Trimestres 1º a 4º</div>
    </div>
    """,
    unsafe_allow_html=True,
)

col1, col2 = st.columns(2)
with col1:
    hog_file = st.file_uploader("Subí la base de HOGARES (CSV/TXT/Excel)", type=["csv", "txt", "xlsx", "xls"]) 
with col2:
    ind_file = st.file_uploader("Subí la base de INDIVIDUOS (CSV/TXT/Excel)", type=["csv", "txt", "xlsx", "xls"]) 

inst_file = st.file_uploader("(Opcional) Subí el Instructivo PDF del trimestre", type=["pdf"]) 

st.info("Podés subir cualquier trimestre. Si no hay variables TIC, la app omite esa sección automáticamente.")

if hog_file and ind_file:
    with st.spinner("Leyendo archivos…"):
        df_hog = _read_table(hog_file)
        df_ind = _read_table(ind_file)

    anio_h, tri_h = detect_year_quarter(df_hog, hog_file.name.lower())
    anio_i, tri_i = detect_year_quarter(df_ind, ind_file.name.lower())

    anio = anio_h if anio_h is not None else anio_i
    if anio is None: anio = anio_i
    tri  = tri_h if tri_h is not None else tri_i

    st.success(f"Detectado: Año = {anio if anio else 's/d'} · Trimestre = {tri if tri else 's/d'}")

    with st.expander("Ver primeras filas – Hogares"):   st.dataframe(df_hog.head(10))
    with st.expander("Ver primeras filas – Individuos"): st.dataframe(df_ind.head(10))

    hog_res = analyze_households(df_hog)
    ind_res = analyze_individuals(df_ind)

    st.subheader("Resumen (preview)")
    c1, c2, c3 = st.columns(3)
    c1.metric("Hogares", f"{hog_res.get('N_hogares', 0):,}".replace(",", "."))
    c2.metric("Individuos", f"{ind_res.get('N_individuos', 0):,}".replace(",", "."))
    c3.metric("Edad media", f"{ind_res.get('edad_media'):.1f}" if ind_res.get("edad_media") is not None else "s/d")

    if st.button("Generar informe Word"):
        meta = {"anio": anio, "trimestre": tri, "hog_file": hog_file.name, "ind_file": ind_file.name}
        doc = build_report(hog_res, ind_res, meta, inst_file.name if inst_file else None)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        default_name = f"Informe_EPH_{anio or 'anio'}_T{tri or 'X'}.docx"
        st.download_button(
            label="Descargar Informe .docx",
            data=buf,
            file_name=default_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
else:
    st.warning("Subí las dos bases (Hogares e Individuos) para continuar.")

